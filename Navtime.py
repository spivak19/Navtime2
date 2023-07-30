import docx
import os
import sqlite3
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from flask import Flask, render_template
import win32api
import win32security
import datetime

def GetAuthor(FilePath):
    file_info = win32api.GetFileAttributes(FilePath)
    owner_sid = win32security.GetFileSecurity(FilePath,
                                              win32security.OWNER_SECURITY_INFORMATION).GetSecurityDescriptorOwner()
    owner_name, domain, _ = win32security.LookupAccountSid(None, owner_sid)
    return owner_name

def extract_data_from_docx(file_path):
    doc = docx.Document(file_path)
    data = []
    for para in doc.paragraphs:
        print(para.text)
        print(para.style.name)

        data.append(para.text)
    return '\n'.join(data)

def GetTitle(FilePath):
    doc = docx.Document(FilePath)
    for para in doc.paragraphs:
        if para.style.name == "Title":
            return para.text


def CreateDB(FolderPath,DbPath):

    conn = sqlite3.connect(DbPath)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY,
            filename TEXT NOT NULL,
            path TEXT NOT NULL,
            author TEXT NOT NULL,
            time TEXT NOT NULL
        )
    ''')
    conn.commit()
    if os.path.exists(FolderPath) and os.path.isdir(FolderPath):
        # Get a list of all files in the folder
        file_list = os.listdir(FolderPath)

        FilesData = []
        # Loop through each file in the folder
        conn = sqlite3.connect(DbPath)
        cursor = conn.cursor()

        for i,file in enumerate(file_list):
            # print(FolderPath + "/" + file)
            FilePath = os.path.abspath(FolderPath + "/" + file)
            filename = GetTitle(FilePath)
            author = GetAuthor(FilePath)
            CTime = os.path.getctime(FilePath)
            CTime = datetime.datetime.fromtimestamp(CTime).strftime("%m/%d/%Y, %H:%M:%S")
            cursor.execute('''
                    INSERT INTO documents (id, filename, path, author, time)
                    VALUES (?, ?, ?, ?, ?)
                ''', (i, filename, FilePath, author, CTime))
            conn.commit()

        conn.close()

def Add2DB(file, FolderPath, DbPath):
    conn = sqlite3.connect(DbPath)
    cursor = conn.cursor()
    FilePath = os.path.abspath(FolderPath + "/" + file)
    filename = GetTitle(FilePath)

    cursor.execute("SELECT * FROM documents ORDER BY id DESC LIMIT 1")
    result = cursor.fetchone()[0]
    author = GetAuthor(FilePath)
    CTime = os.path.getctime(FilePath)
    CTime =  datetime.datetime.fromtimestamp(CTime).strftime("%m/%d/%Y, %H:%M:%S")
    cursor.execute('''
                        INSERT INTO documents (id, filename, path, author, time)
                        VALUES (?, ?, ?, ?, ?)
                    ''', (result+1, filename, FilePath, author, CTime))
    conn.commit()
    conn.close()

def read_data_from_database(DbPath):
    # Connect to the SQLite database
    conn = sqlite3.connect(DbPath)
    cursor = conn.cursor()

    # Execute an SQL query to read data from the database
    cursor.execute('SELECT * FROM documents')

    # Fetch all rows from the query result
    rows = cursor.fetchall()
    # Close the database connection
    conn.close()

    return rows

def CheckAdded(DbPath):
    conn = sqlite3.connect(DbPath)
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM documents ORDER BY id DESC LIMIT 1")
    result = cursor.fetchone()[0]
    conn.close()



class MyHandler(FileSystemEventHandler):
    def __init__(self):
        self.FolderPath = "Docs"
        self.DbPath = "DataBase.db"

    def on_created(self, event):
        if event.is_directory:
            print(f"Directory created: {event.src_path}")
        else:
            time.sleep(2)
            Add2DB(os.path.basename(event.src_path),self.FolderPath,self.DbPath)
            CheckAdded(DbPath)
            print(f"File created: {event.src_path}")

    def on_deleted(self, event):
        if event.is_directory:
            print(f"Directory deleted: {event.src_path}")
        else:
            print(f"File deleted: {event.src_path}")

    def on_modified(self, event):
        if event.is_directory:
            print(f"Directory modified: {event.src_path}")
        else:
            print(f"File modified: {event.src_path}")


if __name__ == '__main__':
    FolderPath = "Docs"
    DbPath = "DataBase.db"
    if not(os.path.exists(DbPath)):
        CreateDB(FolderPath,DbPath)

    event_handler = MyHandler()
    observer = Observer()
    observer.schedule(event_handler, FolderPath, recursive=True)
    print(f"Listening to changes in folder: {FolderPath}")
    observer.start()

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()

    observer.join()








