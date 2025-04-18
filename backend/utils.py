# backend/utils.py
import os
from datetime import datetime
import getpass
from docx import Document
import csv
from pyluach import dates, hebrewcal
from database import SessionLocal
from database import Document as Doc
from sqlalchemy import or_, func

# Define directories relative to this file
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')
OUTPUT_DIR = os.path.join(os.getcwd(), 'files')

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def create_document(template_name, additional_keywords="", new_filename=None, class_text = "סודי"):
    """
    Opens a DOCX template, replaces the placeholder [date] with the current date,
    and the [class] placeholder with class_text, then saves the new document using the provided new_filename (if given).
    """
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    doc = Document(template_path)
    # Format the current date as yyyy-mm-dd (display in the document)
    current_date_display = datetime.now().strftime('%Y-%m-%d')
    
    csv_path = os.path.join(os.path.dirname(__file__), 'user_abbreviations.csv')
    username = getpass.getuser()

    HebrewName = ''
    rank = ''
    title =''
    with open(csv_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            # Compare usernames case-insensitively.
            if row and row[0].strip().lower() == username.lower():
                HebrewName = row[3].strip() + " " + row[4].strip()
                rank = row[6].strip()
                title = row[7].strip()
    
    for para in doc.paragraphs:
        if '[date]' in para.text:
            para.text = para.text.replace('[date]', current_date_display)
        if '[simuchin]' in para.text:
            para.text = para.text.replace('[simuchin]', new_filename)
        if '[doctitle]' in para.text:
            para.text = para.text.replace('[doctitle]', additional_keywords)
        if '[Hebrew date]' in para.text:
            todayHebrew = dates.HebrewDate.today()
            Hdate = str(todayHebrew.hebrew_date_string())
            para.text = para.text.replace('[Hebrew date]', Hdate)
        if '[Hebrew name]' in para.text:
            para.text = para.text.replace('[Hebrew name]', HebrewName)
            
        if '[rank]' in para.text:
            para.text = para.text.replace('[rank]', rank)
        if '[title]' in para.text:
            para.text = para.text.replace('[title]', title)

    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            if '[class]' in para.text:
                para.text = para.text.replace('[class]', class_text)
    # If new_filename is provided, use it (append .docx); otherwise, fall back to default naming.
    
    if new_filename:
        file_name = new_filename + ".docx"
    else:
        # Fallback default naming if needed
        file_name = "NV-" + datetime.now().strftime('%Y%m%d%H%M%S') + ".docx"
    output_path = os.path.join(OUTPUT_DIR, file_name)
    
    doc.save(output_path)
    
    return {
        "file_name": file_name,
        "file_path": output_path,
        "created_at": current_date_display,
        "user": getpass.getuser(),
        "keywords": additional_keywords
    }

def FilenameGen():

    username = getpass.getuser()

    # Read the abbreviation from user_abbreviations.csv.
    # The CSV should have rows in the format: username,abbreviation
    abbrev = None
    csv_path = os.path.join(os.path.dirname(__file__), 'user_abbreviations.csv')
    with open(csv_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            # Compare usernames case-insensitively.
            if row and row[0].strip().lower() == username.lower():
                abbrev = row[5].strip()
                break
    if not abbrev:
        abbrev = "XX"  # Use a default abbreviation if none found.
    today_str = datetime.now().strftime("%d%m%y")

    # Count documents for this user created today.
    db = SessionLocal()
    # We assume that Document.created_at stores a string or datetime;
    # in our utils we saved current_date_display as yyyy-mm-dd.
    # To count, we can filter based on the date prefix if stored as a string.
    # Alternatively, if created_at is a datetime, we can use date comparisons.
    # Here, we assume created_at is a string in "yyyy-mm-dd" format,
    # and we compare the day, month and year (adjust accordingly if needed).
    today_date = datetime.now().date()
    count = db.query(Doc).filter(
        Doc.user == username,
        func.date(Doc.created_at) == today_date
    ).count()
    db.close()
    doc_count = count + 1
    # Build the file name as: NV-{abbrev}-{ddmmyy}-{doc_count}

    print("name generation went well")
    return f"NV-{abbrev}-{today_str}-{doc_count}"