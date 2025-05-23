# backend/server.py
import os
import subprocess
import csv
import getpass
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS 
from werkzeug.utils import secure_filename
from sqlalchemy import or_, func
from docx import Document as DocxDocument

# For Compiling
# from backend.database import SessionLocal, init_db, Document
# from backend.utils import create_document, TEMPLATE_DIR, OUTPUT_DIR


## Before Compiling
from database import SessionLocal, init_db, Document
from utils import create_document, FilenameGen, TEMPLATE_DIR, OUTPUT_DIR

NAME_TO_USER = {}
USER_TO_NAME = {}
csv_path = os.path.join(os.path.dirname(__file__), 'user_abbreviations.csv')

with open(csv_path, newline='', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        if len(row) >= 1:
            username = row['id'].strip()
            name = row['Hebrew First Name'].strip() + " " + row['Hebrew Last Name'].strip()
            abbr = row['Initials'].strip()

            NAME_TO_USER[name] = username
            USER_TO_NAME[username] = name

app = Flask(__name__, static_folder=os.path.join(os.path.dirname(__file__), '../frontend/build'), static_url_path='/')
CORS(app)  # Enable CORS for development

# Initialize the database tables.
init_db()

# Admin password
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'admin123')

@app.route('/api/templates', methods=['GET'])
def list_templates():
    try:
        templates = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith('.docx')]
        return jsonify({"templates": templates})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/create-document', methods=['POST'])
def create_doc():
    data = request.get_json()
    template_name = data.get('template')
    keywords = data.get('keywords', '')
    class_val = data.get('class','סודי')

    if not template_name:
        return jsonify({"error": "Template name is required"}), 400

    try:
        # Retrieve the current logged-in user.
        username = getpass.getuser()
        
        new_filename = FilenameGen()
        # Create the document by calling the updated function.

        print(f'template name: {template_name}, keywords: {keywords}, new filename: {new_filename}')
        
        doc_metadata = create_document(template_name, keywords, new_filename, class_val)
        db = SessionLocal()

        # Save the document metadata in the database.
        new_doc = Document(
            file_name=doc_metadata['file_name'],
            file_path=doc_metadata['file_path'],
            created_at=datetime.now(),  # Storing current datetime.
            user=username,
            keywords=doc_metadata['keywords']
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        db.close()

        return jsonify({"message": "Document created", "document": {
            "id": new_doc.id,
            "file_name": new_doc.file_name,
            "file_path": new_doc.file_path,
            "created_at": new_doc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            "user": new_doc.user,
            "keywords": new_doc.keywords
        }})
    except Exception as e:
        print("error here")
        return jsonify({"error": str(e)}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    
    sort_by = request.args.get('sort_by', 'created_at')
    search = request.args.get('search', '')

    db = SessionLocal()
    query = db.query(Document)
    
    if search:
        matched_users = [
            user for abbr, user in NAME_TO_USER.items()
            if search.lower() in abbr.lower()
        ]
        ors = []
        if matched_users:
            ors = [Document.user == u for u in matched_users]
        query = query.filter(
            or_(
                Document.file_name.like(f"%{search}%"),
                Document.keywords.like(f"%{search}%"),
                *ors,
                Document.user.like(f"%{search}%")
            )
        )

    if sort_by in ['created_at', 'user', 'file_name']:
        query = query.order_by(getattr(Document, sort_by).desc())
    
    documents = query.all()
    db.close()

    docs = []
    for doc in documents:
        # Format created_at if it's a datetime; adjust if stored differently.
        docs.append({
            "id": doc.id,
            "file_name": doc.file_name,
            "file_path": doc.file_path,
            "created_at": doc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            "user": USER_TO_NAME[doc.user],
            "keywords": doc.keywords
        })

    return jsonify({"documents": docs})

@app.route('/api/open-file/<path:filename>', methods=['GET'])
def open_file(filename):
    try:
        return send_from_directory(OUTPUT_DIR, filename, as_attachment=False)
    except Exception as e:
        return jsonify({"error": str(e)}), 404
    

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve(path):
    if path != "" and os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    return send_from_directory(app.static_folder, 'index.html')

# open the links using python method instead of a js one
@app.route('/api/launch-file/<path:filename>', methods=['GET', 'POST'])
def launch_file(filename):
    """
    Tell the OS to open the given file (in OUTPUT_DIR) with its default application.
    """
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404

    try:
        if os.name == 'nt':  # Windows
            os.startfile(file_path)
        elif sys.platform == 'darwin':  # macOS
            subprocess.Popen(['open', file_path])
        else:  # Linux variants
            subprocess.Popen(['xdg-open', file_path])
        return jsonify({"message": "File launched"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/add-file', methods=['POST'])
def add_file():
    """
    Accept a file upload, save it into OUTPUT_DIR, then record it in the DB.
    """
    print("FILES: ",request.files)
    print("FORM: ", request.form.to_dict())

    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    f = request.files['file']
    if f.filename == '':
        return jsonify({"error": "No selected file"}), 400

    # Sanitize filename
    filename = FilenameGen()
    output_path = os.path.join(OUTPUT_DIR, filename)
    # Grab Keywords
    keywords = request.form.get('keywords', '').strip()
    print(keywords)
    # Save the file
    try:
        f.save(output_path)
    except Exception as e:
        return jsonify({"error": f"Failed to save file: {e}"}), 500

    # Record in the database
    try:
        db = SessionLocal()
        new_doc = Document(
            file_name=filename,
            file_path=output_path,
            created_at=datetime.now(),
            user=getpass.getuser(),
            keywords=keywords  # or pull from request.form.get('keywords', '')
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        db.close()
        return jsonify({
            "message": "File added",
            "document": {
                "id": new_doc.id,
                "file_name": new_doc.file_name,
                "file_path": new_doc.file_path,
                "created_at": new_doc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                "user": new_doc.user,
                "keywords": new_doc.keywords
            }
        }), 201
    except Exception as e:
        return jsonify({"error": f"DB error: {e}"}), 500
    

@app.route('/api/import-directory', methods=['POST'])
def import_directory():
    """
    Walks the given directory and adds all .docx files to the DB,
    extracting 'keywords' from each file's core properties.
    Expects JSON:
      {
        "path": "C:/full/path/to/dir",
        "password": "admin123"
      }
    """
    data = request.get_json() or {}
    # 1) Auth check
    if data.get('password') != ADMIN_PASSWORD:
        return jsonify({"error": "Unauthorized"}), 401

    # 2) Path check
    dir_path = data.get('path', '').strip()
    if not dir_path or not os.path.isdir(dir_path):
        return jsonify({"error": "Invalid or missing directory path"}), 400

    db = SessionLocal()
    imported = []
    try:
        for root, _, files in os.walk(dir_path):
            for fname in files:
                if not fname.lower().endswith('.docx'):
                    continue

                full_path = os.path.join(root, fname)
                # Skip if already in DB
                if db.query(Document).filter(Document.file_path == full_path).first():
                    continue

                # 3) Extract keywords and user from the .docx core properties
                try:
                    docx = DocxDocument(full_path)
                    for par in docx.paragraphs:
                        if "הנדון:" in par.text:
                            kw = par.text
                            kw = kw.replace("הנדון:","")
                            break
                    author = docx.core_properties.author
                    print(author)
                except Exception as e:
                    # If the file is corrupted or unreadable, skip it
                    print(f"Failed to read {full_path}: {e}")
                    kw = ''
                    author = getpass.getuser()

                # 4) Use file's last‐modified time as created_at
                ts = os.path.getmtime(full_path)
                created_dt = datetime.fromtimestamp(ts)

                # 5) Insert into DB
                new_doc = Document(
                    file_name=fname,
                    file_path=full_path,
                    created_at=created_dt,
                    user=author,
                    keywords=kw
                )
                db.add(new_doc)
                db.flush()  # so new_doc.id is set
                imported.append({
                    "id": new_doc.id,
                    "file_name": fname,
                    "file_path": full_path,
                    "keywords": kw
                })

        db.commit()
        return jsonify({"imported": imported}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

#Filter search by user
@app.route('/api/users', methods=['GET'])
def get_users():
    """
    Return a JSON list of all distinct user abbreviations in the documents table.
    """
    db = SessionLocal()
    # get raw usernames
    raw_users = [row[0] for row in db.query(Document.user).distinct().all()]
    db.close()

    # map to abbreviations (fall back to raw username)
    abbrs = [USER_TO_NAME.get(u, u) for u in raw_users]

    # sort alphabetically, and return
    abbrs.sort()
    print(abbrs)
    return jsonify({"users": abbrs})


if __name__ == '__main__':
    app.run(debug=True, port=5000)
